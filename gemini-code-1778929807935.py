import streamlit as st
import pandas as pd
import datetime
from docx import Document
from io import BytesIO

# ==========================================
# CẤU HÌNH TRANG & KHỞI TẠO DỮ LIỆU (MOCK DB)
# ==========================================
st.set_page_config(page_title="Quản lý KHTN - TH&THCS Nam Thượng", layout="wide")

if 'chemicals' not in st.session_state:
    st.session_state.chemicals = pd.DataFrame({
        'Mã vật tư': ['HC01', 'HC02', 'VL01', 'SH01'],
        'Tên vật tư': ['Axit Sunfuric (H2SO4)', 'Natri Hidroxit (NaOH)', 'Bộ Khúc xạ ánh sáng', 'Tiêu bản tế bào'],
        'Phân môn': ['Hóa học', 'Hóa học', 'Vật lý', 'Sinh học'],
        'Hạn sử dụng': [datetime.date(2026, 6, 15), datetime.date(2026, 4, 10), None, datetime.date(2027, 1, 1)],
        'Tình trạng': ['Tốt', 'Sắp hết hạn', 'Tốt', 'Tốt']
    })

if 'evaluations' not in st.session_state:
    st.session_state.evaluations = []

# ==========================================
# GIAO DIỆN CHÍNH & ĐIỀU HƯỚNG
# ==========================================
st.sidebar.title("HỆ THỐNG QUẢN LÝ KHTN")
st.sidebar.markdown("**Trường TH&THCS Nam Thượng**")

# Phân quyền người dùng (Mô phỏng đăng nhập)
role = st.sidebar.selectbox("Đăng nhập với vai trò:", 
                            ["Giáo viên bộ môn", "Tổ trưởng chuyên môn", "Phó Hiệu trưởng", "Hiệu trưởng"])

menu = st.sidebar.radio("Chức năng chính", ["Trang chủ & Cảnh báo", "Đăng ký thiết bị", "Đánh giá chuyên môn", "Xuất báo cáo (.docx)"])

# ==========================================
# MODULE 1: TRANG CHỦ & CẢNH BÁO
# ==========================================
if menu == "Trang chủ & Cảnh báo":
    st.header("Bảng điều khiển (Dashboard)")
    
    col1, col2, col3 = st.columns(3)
    col1.metric("Tổng thiết bị/Hóa chất", len(st.session_state.chemicals))
    col2.metric("Số tiết đã đánh giá", len(st.session_state.evaluations))
    
    st.subheader("Cảnh báo an toàn hóa chất (Sắp hết hạn dưới 30 ngày)")
    today = datetime.date.today()
    
    # Lọc hóa chất có hạn sử dụng và sắp hết hạn
    df = st.session_state.chemicals.copy()
    df_exp = df.dropna(subset=['Hạn sử dụng'])
    df_warning = df_exp[(df_exp['Hạn sử dụng'] - today).dt.days <= 30]
    
    if not df_warning.empty:
        st.error("CẢNH BÁO: Phát hiện hóa chất/tiêu bản sắp hết hạn hoặc đã hết hạn!")
        st.dataframe(df_warning)
    else:
        st.success("Tất cả hóa chất và tiêu bản đều trong thời hạn sử dụng an toàn.")
        
    st.subheader("Danh mục vật tư hiện có")
    st.dataframe(st.session_state.chemicals)

# ==========================================
# MODULE 2: ĐĂNG KÝ THIẾT BỊ
# ==========================================
elif menu == "Đăng ký thiết bị":
    st.header("Đăng ký sử dụng phòng bộ môn & Thiết bị")
    with st.form("booking_form"):
        col1, col2 = st.columns(2)
        with col1:
            date = st.date_input("Ngày thực hành")
            period = st.selectbox("Tiết học", [1, 2, 3, 4, 5])
            subject = st.selectbox("Phân môn", ["Vật lý", "Hóa học", "Sinh học"])
        with col2:
            equipment = st.multiselect("Chọn thiết bị/Hóa chất cần mượn", st.session_state.chemicals['Tên vật tư'].tolist())
            note = st.text_input("Ghi chú (Tình trạng ban đầu)")
            
        submitted = st.form_submit_button("Xác nhận đăng ký")
        if submitted:
            # Trong thực tế sẽ có thuật toán check trùng lịch ở đây
            st.success("Đăng ký thành công! Hệ thống đã ghi nhận lịch của bạn.")

# ==========================================
# MODULE 3: ĐÁNH GIÁ CHUYÊN MÔN THEO RUBRIC
# ==========================================
elif menu == "Đánh giá chuyên môn":
    st.header("Đánh giá năng lực tổ chức thực hành KHTN")
    
    # Luồng đánh giá phân cấp
    eval_target = st.text_input("Họ và tên người được đánh giá:")
    
    if role == "Hiệu trưởng":
        st.info("Luồng đánh giá: Hiệu trưởng đánh giá Tổ trưởng / Phó Hiệu trưởng")
    elif role == "Phó Hiệu trưởng":
        st.info("Luồng đánh giá: Phó Hiệu trưởng đánh giá Giáo viên / Tổ trưởng")
    elif role == "Tổ trưởng chuyên môn":
        st.info("Luồng đánh giá: Tổ trưởng đánh giá Giáo viên bộ môn")
    else:
        st.info("Luồng tự đánh giá của Giáo viên")

    st.markdown("### Chấm điểm Rubric")
    c1 = st.slider("1. Công tác chuẩn bị thiết bị, vật tư (Tối đa 20đ)", 0, 20, 15)
    c2 = st.slider("2. Đảm bảo quy tắc an toàn PTN (Tối đa 30đ)", 0, 30, 25)
    c3 = st.slider("3. Tổ chức và hướng dẫn HS thao tác (Tối đa 30đ)", 0, 30, 25)
    c4 = st.slider("4. Đánh giá kết quả & Liên hệ PISA (Tối đa 20đ)", 0, 20, 15)
    
    total = c1 + c2 + c3 + c4
    st.markdown(f"**Tổng điểm:** {total}/100")
    
    if total >= 85: rank = "Tốt"
    elif total >= 65: rank = "Khá"
    elif total >= 50: rank = "Đạt"
    else: rank = "Chưa đạt"
    
    comment = st.text_area("Nhận xét ưu điểm & Tồn tại:")
    
    if st.button("Lưu đánh giá"):
        if eval_target:
            record = {
                "Người được đánh giá": eval_target,
                "Người đánh giá": role,
                "Tổng điểm": total,
                "Xếp loại": rank,
                "Nhận xét": comment,
                "Ngày": datetime.date.today().strftime("%d/%m/%Y")
            }
            st.session_state.evaluations.append(record)
            st.success("Đã lưu hồ sơ đánh giá vào hệ thống!")
        else:
            st.warning("Vui lòng nhập tên người được đánh giá.")

# ==========================================
# MODULE 4: XUẤT BÁO CÁO TỰ ĐỘNG (.DOCX)
# ==========================================
elif menu == "Xuất báo cáo (.docx)":
    st.header("Kết xuất hồ sơ minh chứng")
    
    if len(st.session_state.evaluations) == 0:
        st.info("Chưa có hồ sơ đánh giá nào trong hệ thống.")
    else:
        # Hiển thị danh sách đánh giá
        df_evals = pd.DataFrame(st.session_state.evaluations)
        st.dataframe(df_evals)
        
        # Chọn một hồ sơ để xuất file
        selected_idx = st.selectbox("Chọn hồ sơ cần xuất báo cáo", range(len(st.session_state.evaluations)), 
                                    format_func=lambda x: f"Hồ sơ: {st.session_state.evaluations[x]['Người được đánh giá']} - {st.session_state.evaluations[x]['Ngày']}")
        
        target_record = st.session_state.evaluations[selected_idx]
        
        # Hàm tạo file Word
        def create_docx(data):
            doc = Document()
            doc.add_heading('PHIẾU ĐÁNH GIÁ NĂNG LỰC THỰC HÀNH', 0)
            doc.add_paragraph(f"Trường: TH&THCS Nam Thượng")
            doc.add_paragraph(f"Ngày đánh giá: {data['Ngày']}")
            doc.add_paragraph(f"Người được đánh giá: {data['Người được đánh giá']}")
            doc.add_paragraph(f"Cấp đánh giá: {data['Người đánh giá']}")
            
            doc.add_heading('Kết quả', level=1)
            doc.add_paragraph(f"- Tổng điểm: {data['Tổng điểm']}/100")
            doc.add_paragraph(f"- Xếp loại: {data['Xếp loại']}")
            
            doc.add_heading('Nhận xét chuyên môn', level=1)
            doc.add_paragraph(data['Nhận xét'])
            
            # Lưu vào bộ nhớ ảo (BytesIO) để tải về
            bio = BytesIO()
            doc.save(bio)
            return bio.getvalue()

        # Nút tải xuống
        docx_file = create_docx(target_record)
        st.download_button(
            label="📄 Tải xuống báo cáo (.docx)",
            data=docx_file,
            file_name=f"Bao_cao_danh_gia_{target_record['Người được đánh giá']}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )