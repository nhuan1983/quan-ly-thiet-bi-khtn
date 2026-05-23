import streamlit as st
import pandas as pd
import datetime
import json
from docx import Document
from io import BytesIO

# ==========================================
# CẤU HÌNH KẾT NỐI GOOGLE SHEETS (DATABASE REAL)
# ==========================================
USE_CLOUD_DB = False
conn = None
sh = None

if "gspread_creds" in st.secrets and "spreadsheet_key" in st.secrets:
    try:
        import gspread
        creds_dict = dict(st.secrets["gspread_creds"])
        if "private_key" in creds_dict:
            creds_dict["private_key"] = creds_dict["private_key"].replace("\\n", "\n")
            
        gc = gspread.service_account_from_dict(creds_dict)
        sh = gc.open_by_key(st.secrets["spreadsheet_key"])
        USE_CLOUD_DB = True
    except Exception as e:
        st.sidebar.error(f"⚠️ Lỗi kết nối Google Sheets: {e}")

# --- CÁC HÀM ĐỌC/GHI DỮ LIỆU THÔNG MINH ---
def load_data(sheet_name, default_df):
    if USE_CLOUD_DB:
        try:
            worksheet = sh.worksheet(sheet_name)
            records = worksheet.get_all_records()
            if not records:
                return default_df
            df = pd.DataFrame(records)
            if sheet_name == 'users' and 'Vai trò' in df.columns:
                df['Vai trò'] = df['Vai trò'].apply(lambda x: [r.strip() for r in str(x).split(',')])
            return df
        except Exception:
            return default_df
    return default_df

def save_data(sheet_name, df_to_save):
    if USE_CLOUD_DB:
        try:
            try:
                worksheet = sh.worksheet(sheet_name)
            except gspread.exceptions.WorksheetNotFound:
                worksheet = sh.add_worksheet(title=sheet_name, rows="100", cols="20")
            
            worksheet.clear()
            df_copy = df_to_save.copy()
            if sheet_name == 'users' and 'Vai trò' in df_copy.columns:
                df_copy['Vai trò'] = df_copy['Vai trò'].apply(lambda x: ", ".join(x) if isinstance(x, list) else x)
            
            for col in df_copy.columns:
                if df_copy[col].dtype == 'object':
                    df_copy[col] = df_copy[col].astype(str)
                    
            worksheet.update([df_copy.columns.values.tolist()] + df_copy.values.tolist())
        except Exception as e:
            st.error(f"Không thể đồng bộ lên Google Sheets: {e}")

# ==========================================
# KHỞI TẠO HOẶC TẢI CƠ SỞ DỮ LIỆU
# ==========================================
default_school = pd.DataFrame([{'ten_truong': 'TH&THCS Nam Thượng', 'don_vi_chu_quan': 'Ủy ban nhân dân xã Hợp Kim', 'nam_hoc': '2025-2026'}])
df_school_db = load_data('school_info', default_school)
if 'school_info' not in st.session_state:
    st.session_state.school_info = df_school_db.iloc[0].to_dict()

st.set_page_config(page_title=f"Quản lý KHTN - {st.session_state.school_info['ten_truong']}", layout="wide")

default_users = pd.DataFrame({
    'Tài khoản': ['admin', 'ht', 'totruong', 'gv01'],
    'Mật khẩu': ['123', '123', '123', '123'],
    'Họ tên': ['Quản trị viên (PHT)', 'Nguyễn Văn A (Hiệu trưởng)', 'Trần Thị B (Tổ trưởng)', 'Lê Văn C (Giáo viên)'],
    'Vai trò': [['Quản trị viên', 'Phó Hiệu trưởng', 'Giáo viên bộ môn'], ['Hiệu trưởng', 'Giáo viên bộ môn'], ['Tổ trưởng chuyên môn', 'Giáo viên bộ môn'], ['Giáo viên bộ môn']]
})
if 'users' not in st.session_state:
    st.session_state.users = load_data('users', default_users)

default_chem = pd.DataFrame({
    'Mã vật tư': ['HC01', 'HC02', 'VL01', 'SH01'],
    'Tên vật tư': ['Axit Sunfuric (H2SO4)', 'Natri Hidroxit (NaOH)', 'Bộ Khúc xạ ánh sáng', 'Tiêu bản tế bào'],
    'Phân môn': ['Hóa học', 'Hóa học', 'Vật lý', 'Sinh học'],
    'Số lượng': [5, 10, 3, 20],
    'Hạn sử dụng': ['2026-06-15', '2026-04-10', 'None', '2027-01-01'],
    'Tình trạng': ['Tốt', 'Sắp hết hạn', 'Tốt', 'Tốt']
})
if 'chemicals' not in st.session_state:
    st.session_state.chemicals = load_data('chemicals', default_chem)

if 'bookings' not in st.session_state:
    st.session_state.bookings = load_data('bookings', pd.DataFrame(columns=['Người đăng ký', 'Ngày', 'Buổi', 'Tiết', 'Lớp', 'Môn', 'Thiết bị']))

if 'evaluations' not in st.session_state:
    df_eval_loaded = load_data('evaluations', pd.DataFrame())
    st.session_state.evaluations = df_eval_loaded.to_dict('records') if not df_eval_loaded.empty else []

if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False
if 'current_user' not in st.session_state:
    st.session_state.current_user = None

# ==========================================
# GIAO DIỆN ĐĂNG NHẬP
# ==========================================
if not st.session_state.logged_in:
    st.markdown("<h1 style='text-align: center; color: #1E88E5;'>HỆ THỐNG QUẢN LÝ KHTN</h1>", unsafe_allow_html=True)
    st.markdown(f"<h3 style='text-align: center;'>Trường {st.session_state.school_info['ten_truong']}</h3>", unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        with st.form("login_form"):
            st.write("Vui lòng đăng nhập để tiếp tục")
            username = st.text_input("Tài khoản")
            password = st.text_input("Mật khẩu", type="password")
            submit_login = st.form_submit_button("Đăng nhập", use_container_width=True)
            
            if submit_login:
                # Ép kiểu chuỗi để sửa lỗi với Google Sheets
                user_match = st.session_state.users[(st.session_state.users['Tài khoản'].astype(str) == str(username)) & (st.session_state.users['Mật khẩu'].astype(str) == str(password))]
                if not user_match.empty:
                    st.session_state.logged_in = True
                    st.session_state.current_user = user_match.iloc[0].to_dict()
                    st.rerun()
                else:
                    st.error("Sai tài khoản hoặc mật khẩu!")
    st.stop()

# ==========================================
# SIDEBAR NAVIGATION
# ==========================================
current_user = st.session_state.current_user
user_roles = current_user['Vai trò']

st.sidebar.title(st.session_state.school_info['ten_truong'])
if USE_CLOUD_DB:
    st.sidebar.caption("🟢 Đã kết nối cơ sở dữ liệu đám mây")
else:
    st.sidebar.caption("🔵 Chế độ cục bộ (Lưu trữ tạm thời)")

st.sidebar.success(f"👤 Chào, {current_user['Họ tên']}")
st.sidebar.markdown("---")
active_role = st.sidebar.selectbox("🔄 Bạn đang làm việc với tư cách là:", user_roles)
st.sidebar.markdown("---")

menu_options = ["Trang chủ & Cảnh báo", "Quản lý Kho (Vật tư)", "Đăng ký thiết bị"]
if active_role in ["Quản trị viên", "Hiệu trưởng", "Phó Hiệu trưởng", "Tổ trưởng chuyên môn"]:
    menu_options.append("Đánh giá chuyên môn")
    menu_options.append("Xuất báo cáo (.docx)")
if active_role == "Quản trị viên":
    menu_options.insert(0, "Quản lý Hệ thống (Admin)")

menu = st.sidebar.radio("📌 Chọn chức năng:", menu_options)

if st.sidebar.button("Đăng xuất"):
    st.session_state.logged_in = False
    st.session_state.current_user = None
    st.rerun()

# ==========================================
# MODULE: QUẢN LÝ HỆ THỐNG (UI TABS MỚI)
# ==========================================
if menu == "Quản lý Hệ thống (Admin)":
    st.header("⚙️ Quản lý Hệ thống & Cấu hình Đơn vị")
    
    # CHIA LÀM 3 TABS GỌN GÀNG
    tab1, tab2, tab3 = st.tabs(["🏫 Cấu hình Trường học", "👥 Quản lý Tài khoản", "📥 Nhập dữ liệu Excel"])
    
    with tab1:
        st.subheader("Cấu hình thông tin Trường học & Chính quyền địa phương")
        with st.form("school_config_form"):
            sc_col1, sc_col2, sc_col3 = st.columns(3)
            edit_ten_truong = sc_col1.text_input("Tên trường", value=st.session_state.school_info['ten_truong'])
            edit_chu_quan = sc_col2.text_input("Đơn vị chủ quản", value=st.session_state.school_info['don_vi_chu_quan'])
            edit_nam_hoc = sc_col3.text_input("Năm học", value=st.session_state.school_info['nam_hoc'])
            if st.form_submit_button("💾 Lưu cấu hình"):
                st.session_state.school_info['ten_truong'] = edit_ten_truong
                st.session_state.school_info['don_vi_chu_quan'] = edit_chu_quan
                st.session_state.school_info['nam_hoc'] = edit_nam_hoc
                save_data('school_info', pd.DataFrame([st.session_state.school_info]))
                st.success("Đã lưu cấu hình!")
                st.rerun()

    with tab2:
        st.subheader("Danh sách tài khoản hiện tại")
        df_display = st.session_state.users.copy()
        df_display['Vai trò'] = df_display['Vai trò'].apply(lambda x: ", ".join(x) if isinstance(x, list) else x)
        st.dataframe(df_display, use_container_width=True)
        
        col_left, col_right = st.columns(2)
        with col_left:
            st.markdown("**➕ Cấp tài khoản mới**")
            with st.form("add_user_form"):
                new_acc = st.text_input("Tên đăng nhập")
                new_pwd = st.text_input("Mật khẩu")
                new_name = st.text_input("Họ và tên")
                new_roles = st.multiselect("Phân quyền", ["Giáo viên bộ môn", "Tổ trưởng chuyên môn", "Phó Hiệu trưởng", "Hiệu trưởng", "Quản trị viên"])
                if st.form_submit_button("Tạo tài khoản"):
                    if new_acc and new_pwd and new_name and new_roles:
                        new_row = pd.DataFrame([{'Tài khoản': new_acc, 'Mật khẩu': new_pwd, 'Họ tên': new_name, 'Vai trò': new_roles}])
                        st.session_state.users = pd.concat([st.session_state.users, new_row], ignore_index=True)
                        save_data('users', st.session_state.users)
                        st.success("Đã tạo!")
                        st.rerun()
        
        with col_right:
            st.markdown("**✏️ Sửa / ❌ Xóa tài khoản**")
            user_list = st.session_state.users['Tài khoản'].tolist()
            selected_user = st.selectbox("Chọn tài khoản:", ["-- Chọn --"] + user_list)
            if selected_user != "-- Chọn --":
                user_data = st.session_state.users[st.session_state.users['Tài khoản'] == selected_user].iloc[0]
                with st.form("edit_delete_user_form"):
                    edit_name = st.text_input("Họ và tên", value=user_data['Họ tên'])
                    edit_pwd = st.text_input("Mật khẩu", value=user_data['Mật khẩu'])
                    current_roles = user_data['Vai trò'] if isinstance(user_data['Vai trò'], list) else [r.strip() for r in str(user_data['Vai trò']).split(',')]
                    edit_roles = st.multiselect("Phân quyền", ["Giáo viên bộ môn", "Tổ trưởng chuyên môn", "Phó Hiệu trưởng", "Hiệu trưởng", "Quản trị viên"], default=current_roles)
                    btn1, btn2 = st.columns(2)
                    if btn1.form_submit_button("💾 Lưu"):
                        idx = st.session_state.users[st.session_state.users['Tài khoản'] == selected_user].index[0]
                        st.session_state.users.at[idx, 'Họ tên'] = edit_name
                        st.session_state.users.at[idx, 'Mật khẩu'] = edit_pwd
                        st.session_state.users.at[idx, 'Vai trò'] = edit_roles
                        save_data('users', st.session_state.users)
                        st.success("Đã lưu!")
                        st.rerun()
                    if btn2.form_submit_button("❌ Xóa"):
                        if selected_user != current_user['Tài khoản']:
                            st.session_state.users = st.session_state.users[st.session_state.users['Tài khoản'] != selected_user].reset_index(drop=True)
                            save_data('users', st.session_state.users)
                            st.success("Đã xóa!")
                            st.rerun()
                        else:
                            st.error("Không thể tự xóa!")

    with tab3:
        st.subheader("Nhập hàng loạt tài khoản từ Excel")
        df_mau_tk = pd.DataFrame({'Tài khoản': ['gv_toan01'], 'Mật khẩu': ['123'], 'Họ tên': ['Trần Thị D'], 'Vai trò': ['Giáo viên bộ môn']})
        output_tk = BytesIO()
        with pd.ExcelWriter(output_tk, engine='openpyxl') as writer:
            df_mau_tk.to_excel(writer, index=False)
        st.download_button(label="⬇️ Tải file Excel mẫu", data=output_tk.getvalue(), file_name="Mau_Tai_Khoan.xlsx")
        
        uploaded_users = st.file_uploader("Kéo thả file Excel vào đây", type=['xlsx', 'xls'])
        if uploaded_users is not None and st.button("Tiến hành nhập"):
            try:
                df_new = pd.read_excel(uploaded_users)
                df_new['Vai trò'] = df_new['Vai trò'].astype(str).apply(lambda x: [r.strip() for r in x.split(',')])
                df_new['Tài khoản'] = df_new['Tài khoản'].astype(str)
                df_new['Mật khẩu'] = df_new['Mật khẩu'].astype(str)
                st.session_state.users = pd.concat([st.session_state.users, df_new], ignore_index=True)
                save_data('users', st.session_state.users)
                st.success("Nhập thành công!")
                st.rerun()
            except Exception as e:
                st.error(f"Lỗi: {e}")

# ==========================================
# MODULE: TRANG CHỦ & CẢNH BÁO
# ==========================================
elif menu == "Trang chủ & Cảnh báo":
    st.header("📊 Bảng điều khiển (Dashboard)")
    col1, col2 = st.columns(2)
    col1.metric("Tổng số thiết bị/Hóa chất", len(st.session_state.chemicals))
    col2.metric("Số lượt đăng ký mượn", len(st.session_state.bookings))
    
    today = datetime.date.today()
    df_chem_check = st.session_state.chemicals.copy()
    df_chem_check['Hạn sử dụng'] = pd.to_datetime(df_chem_check['Hạn sử dụng'], errors='coerce').dt.date
    df_exp = df_chem_check.dropna(subset=['Hạn sử dụng'])
    df_warning = df_exp[(df_exp['Hạn sử dụng'] - today).dt.days <= 30]
    
    st.subheader("⚠️ Cảnh báo an toàn")
    if not df_warning.empty:
        st.error("Phát hiện vật tư sắp hết hạn!")
        st.dataframe(df_warning, use_container_width=True)
    else:
        st.success("Tất cả hóa chất an toàn.")

# ==========================================
# MODULE: QUẢN LÝ KHO (UI TABS MỚI)
# ==========================================
elif menu == "Quản lý Kho (Vật tư)":
    st.header("📦 Quản lý Kho Thiết bị & Hóa chất")
    
    if active_role in ["Quản trị viên", "Tổ trưởng chuyên môn", "Phó Hiệu trưởng", "Hiệu trưởng"]:
        tab1, tab2, tab3 = st.tabs(["📦 Xem Danh mục", "➕ Nhập/Thêm mới", "⚙️ Chỉnh sửa"])
    else:
        tab1, = st.tabs(["📦 Xem Danh mục"])
        tab2 = tab3 = None

    with tab1:
        st.dataframe(st.session_state.chemicals, use_container_width=True)
    
    if tab2:
        with tab2:
            col_add, col_import = st.columns(2)
            with col_add:
                st.markdown("**Bổ sung thủ công**")
                with st.form("add_chem_form"):
                    ma_vt = st.text_input("Mã vật tư")
                    ten_vt = st.text_input("Tên vật tư")
                    phan_mon = st.selectbox("Phân môn", ["Vật lý", "Hóa học", "Sinh học"])
                    so_luong = st.number_input("Số lượng", min_value=1, value=1)
                    han_su_dung = st.date_input("Hạn sử dụng", value=None)
                    tinh_trang = st.selectbox("Tình trạng", ["Tốt", "Cần sửa chữa", "Đang đặt mua"])
                    if st.form_submit_button("Lưu vào kho") and ma_vt and ten_vt:
                        new_item = pd.DataFrame([{'Mã vật tư': ma_vt, 'Tên vật tư': ten_vt, 'Phân môn': phan_mon, 'Số lượng': int(so_luong), 'Hạn sử dụng': str(han_su_dung), 'Tình trạng': tinh_trang}])
                        st.session_state.chemicals = pd.concat([st.session_state.chemicals, new_item], ignore_index=True)
                        save_data('chemicals', st.session_state.chemicals)
                        st.success("Đã thêm!")
                        st.rerun()
            with col_import:
                st.markdown("**Nhập từ Excel**")
                df_mau_vt = pd.DataFrame({'Mã vật tư': ['VL02'], 'Tên vật tư': ['Ampe kế'], 'Phân môn': ['Vật lý'], 'Số lượng': [15], 'Hạn sử dụng': [''], 'Tình trạng': ['Tốt']})
                output_vt = BytesIO()
                with pd.ExcelWriter(output_vt, engine='openpyxl') as writer:
                    df_mau_vt.to_excel(writer, index=False)
                st.download_button("⬇️ Tải file mẫu", data=output_vt.getvalue(), file_name="Mau_Vat_Tu.xlsx")
                
                uploaded_chem = st.file_uploader("Kéo thả file đã điền", type=['xlsx', 'xls'])
                if uploaded_chem is not None and st.button("Tiến hành nhập"):
                    try:
                        df_new = pd.read_excel(uploaded_chem)
                        st.session_state.chemicals = pd.concat([st.session_state.chemicals, df_new], ignore_index=True)
                        save_data('chemicals', st.session_state.chemicals)
                        st.success("Nhập thành công!")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Lỗi: {e}")
                        
    if tab3:
        with tab3:
            st.markdown("**Chỉnh sửa hoặc Xóa**")
            item_list = st.session_state.chemicals['Mã vật tư'].tolist()
            selected_item_code = st.selectbox("Chọn Mã vật tư:", ["-- Chọn --"] + item_list)
            if selected_item_code != "-- Chọn --":
                item_data = st.session_state.chemicals[st.session_state.chemicals['Mã vật tư'] == selected_item_code].iloc[0]
                with st.form("edit_delete_item_form"):
                    edit_ten = st.text_input("Tên vật tư", value=item_data['Tên vật tư'])
                    col_e1, col_e2 = st.columns(2)
                    edit_mon = col_e1.selectbox("Phân môn", ["Vật lý", "Hóa học", "Sinh học"], index=["Vật lý", "Hóa học", "Sinh học"].index(item_data['Phân môn']))
                    default_sl = int(item_data['Số lượng']) if 'Số lượng' in item_data and pd.notnull(item_data['Số lượng']) else 1
                    edit_sl = col_e2.number_input("Số lượng", min_value=0, value=default_sl)
                    edit_tt = st.selectbox("Tình trạng", ["Tốt", "Cần sửa chữa", "Đang đặt mua"], index=["Tốt", "Cần sửa chữa", "Đang đặt mua"].index(item_data['Tình trạng']))
                    
                    ic_col1, ic_col2 = st.columns(2)
                    if ic_col1.form_submit_button("💾 Lưu"):
                        idx = st.session_state.chemicals[st.session_state.chemicals['Mã vật tư'] == selected_item_code].index[0]
                        st.session_state.chemicals.at[idx, 'Tên vật tư'] = edit_ten
                        st.session_state.chemicals.at[idx, 'Phân môn'] = edit_mon
                        st.session_state.chemicals.at[idx, 'Số lượng'] = int(edit_sl)
                        st.session_state.chemicals.at[idx, 'Tình trạng'] = edit_tt
                        save_data('chemicals', st.session_state.chemicals)
                        st.success("Đã lưu!")
                        st.rerun()
                    if ic_col2.form_submit_button("❌ Xóa"):
                        st.session_state.chemicals = st.session_state.chemicals[st.session_state.chemicals['Mã vật tư'] != selected_item_code].reset_index(drop=True)
                        save_data('chemicals', st.session_state.chemicals)
                        st.success("Đã xóa!")
                        st.rerun()

# ==========================================
# MODULE: ĐĂNG KÝ THIẾT BỊ (UI TABS MỚI)
# ==========================================
elif menu == "Đăng ký thiết bị":
    st.header("📝 Đăng ký sử dụng phòng bộ môn")
    
    tab1, tab2 = st.tabs(["📅 Lịch toàn trường", "📝 Tạo đăng ký mới"])
    
    with tab1:
        if not st.session_state.bookings.empty:
            st.dataframe(st.session_state.bookings, use_container_width=True)
        else:
            st.info("Chưa có lịch đăng ký nào.")
            
    with tab2:
        with st.form("booking_form"):
            col1, col2 = st.columns(2)
            with col1:
                date = st.date_input("Ngày dạy")
                buoi = st.selectbox("Buổi dạy", ["Sáng", "Chiều"])
                period = st.selectbox("Tiết học", [1, 2, 3, 4, 5])
                lop = st.text_input("Lớp (VD: 9A)")
                subject = st.selectbox("Phân môn", ["Vật lý", "Hóa học", "Sinh học"])
            with col2:
                equipment = st.multiselect("Chọn thiết bị", st.session_state.chemicals['Tên vật tư'].tolist())
                
            if st.form_submit_button("Xác nhận đăng ký"):
                if lop:
                    new_book = pd.DataFrame([{'Người đăng ký': current_user['Họ tên'], 'Ngày': str(date), 'Buổi': buoi, 'Tiết': period, 'Lớp': lop, 'Môn': subject, 'Thiết bị': ", ".join(equipment)}])
                    st.session_state.bookings = pd.concat([st.session_state.bookings, new_book], ignore_index=True)
                    save_data('bookings', st.session_state.bookings)
                    st.success("Lịch đã được lưu!")
                    st.rerun()

# ==========================================
# MODULE: ĐÁNH GIÁ CHUYÊN MÔN
# ==========================================
elif menu == "Đánh giá chuyên môn":
    st.header("📋 Đánh giá năng lực thực hành")
    list_gv = st.session_state.users['Họ tên'].tolist()
    target_gv = st.selectbox("1. Chọn Giáo viên:", ["-- Chọn --"] + list_gv)
    
    if target_gv != "-- Chọn --":
        gv_bookings = st.session_state.bookings[st.session_state.bookings['Người đăng ký'] == target_gv]
        if gv_bookings.empty:
            st.warning("Chưa đăng ký tiết dạy.")
        else:
            booking_options = [f"Ngày {row['Ngày']} - Buổi {row['Buổi']} - Tiết {row['Tiết']} - Lớp {row['Lớp']} - Môn {row['Môn']}" for idx, row in gv_bookings.iterrows()]
            target_tiet = st.selectbox("2. Chọn Tiết dự giờ:", booking_options)
            
            st.markdown("### 3. Rubric chuyên môn")
            c1 = st.slider("1. Chuẩn bị thiết bị vật tư (20đ)", 0, 20, 15)
            c2 = st.slider("2. Đảm bảo an toàn PTN (30đ)", 0, 30, 25)
            c3 = st.slider("3. Hướng dẫn HS (30đ)", 0, 30, 25)
            c4 = st.slider("4. Đánh giá & Liên hệ PISA (20đ)", 0, 20, 15)
            
            total = c1 + c2 + c3 + c4
            st.markdown(f"**Tổng điểm:** {total}/100")
            rank = "Tốt" if total >= 85 else "Khá" if total >= 65 else "Đạt" if total >= 50 else "Chưa đạt"
            comment = st.text_area("Nhận xét:")
            
            if st.button("Lưu hồ sơ"):
                record = {
                    "Người được đánh giá": target_gv, "Tiết dạy": target_tiet, "Người đánh giá": current_user['Họ tên'],
                    "Chức vụ người đánh giá": active_role, "Tổng điểm": total, "Xếp loại": rank, "Nhận xét": comment,
                    "Ngày chấm": datetime.date.today().strftime("%d/%m/%Y")
                }
                st.session_state.evaluations.append(record)
                save_data('evaluations', pd.DataFrame(st.session_state.evaluations))
                st.success("Đã đồng bộ lên Google Sheets!")

# ==========================================
# MODULE: XUẤT BÁO CÁO TỰ ĐỘNG (.DOCX)
# ==========================================
elif menu == "Xuất báo cáo (.docx)":
    st.header("🖨️ Kết xuất minh chứng")
    if len(st.session_state.evaluations) == 0:
        st.info("Chưa có hồ sơ.")
    else:
        st.dataframe(pd.DataFrame(st.session_state.evaluations), use_container_width=True)
        selected_idx = st.selectbox("Chọn hồ sơ:", range(len(st.session_state.evaluations)), format_func=lambda x: f"Đánh giá: {st.session_state.evaluations[x]['Người được đánh giá']} ({st.session_state.evaluations[x]['Tiết dạy']})")
        target_record = st.session_state.evaluations[selected_idx]
        
        def create_docx(data, school_info):
            doc = Document()
            doc.add_heading(school_info['don_vi_chu_quan'].upper(), 1)
            doc.add_heading(f"TRƯỜNG {school_info['ten_truong'].upper()}", 2)
            doc.add_paragraph(f"Năm học: {school_info['nam_hoc']}")
            doc.add_paragraph("-----------------------------------")
            doc.add_heading('PHIẾU ĐÁNH GIÁ NĂNG LỰC THỰC HÀNH', 0)
            doc.add_paragraph(f"Ngày lập: {data['Ngày chấm']}")
            doc.add_paragraph(f"Người đánh giá: {data['Người đánh giá']} ({data['Chức vụ người đánh giá']})")
            doc.add_heading('Thông tin tiết dạy', level=1)
            doc.add_paragraph(f"- Giáo viên: {data['Người được đánh giá']}")
            doc.add_paragraph(f"- Lịch học: {data['Tiết dạy']}")
            doc.add_heading('Kết quả', level=1)
            doc.add_paragraph(f"- Điểm: {data['Tổng điểm']}/100")
            doc.add_paragraph(f"- Xếp loại: {data['Xếp loại']}")
            doc.add_heading('Nhận xét', level=1)
            doc.add_paragraph(data['Nhận xét'])
            bio = BytesIO()
            doc.save(bio)
            return bio.getvalue()

        st.download_button("📄 Tải xuống văn bản (.docx)", data=create_docx(target_record, st.session_state.school_info), file_name=f"PhieuDanhGia_{target_record['Người được đánh giá']}.docx")
